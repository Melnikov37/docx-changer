"""
Скрипт для конвертации одинарных фигурных скобок {variable}
в двойные {{variable}} для совместимости с Jinja2
+ замена пробелов в именах переменных на подчеркивания
"""
import re
from docx import Document
import os


def normalize_variable_name(name):
    """Нормализация имени переменной - замена пробелов на подчеркивания"""
    # Заменяем пробелы на подчеркивания
    normalized = re.sub(r'\s+', '_', name)
    # Убираем лишние подчеркивания в начале и конце
    normalized = normalized.strip('_')
    # Заменяем множественные подчеркивания на одинарные
    normalized = re.sub(r'_+', '_', normalized)
    return normalized


def process_paragraph(paragraph):
    """Обработка параграфа - замена {var} на {{var}} и нормализация имен"""
    full_text = paragraph.text

    # Паттерн: {слово} но НЕ {{слово}} и НЕ }}
    pattern = r'(?<!\{)\{([^{}]+)\}(?!\})'

    # Проверяем, есть ли что заменять
    if not re.search(pattern, full_text):
        return False, {}

    # Словарь для отслеживания замен
    replacements = {}

    def replace_func(match):
        """Функция замены с нормализацией"""
        original_name = match.group(1)
        normalized_name = normalize_variable_name(original_name)
        replacements[original_name] = normalized_name
        return '{{' + normalized_name + '}}'

    # Заменяем одинарные скобки на двойные с нормализацией
    new_text = re.sub(pattern, replace_func, full_text)

    # Очищаем все runs в параграфе
    for run in paragraph.runs:
        run.text = ''

    # Добавляем новый текст в первый run (или создаем новый)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)

    return True, replacements


def process_table(table):
    """Обработка таблицы"""
    changed = False
    all_replacements = {}

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                was_changed, replacements = process_paragraph(paragraph)
                if was_changed:
                    changed = True
                    all_replacements.update(replacements)

    return changed, all_replacements


def convert_document(input_path, output_path=None):
    """Конвертация DOCX файла"""
    if output_path is None:
        # Создаем имя выходного файла
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_converted{ext}"

    print(f"Обработка: {input_path}")

    try:
        # Загружаем документ
        doc = Document(input_path)

        changes_count = 0
        all_replacements = {}

        # Обрабатываем параграфы
        for paragraph in doc.paragraphs:
            was_changed, replacements = process_paragraph(paragraph)
            if was_changed:
                changes_count += 1
                all_replacements.update(replacements)

        # Обрабатываем таблицы
        for table in doc.tables:
            was_changed, replacements = process_table(table)
            if was_changed:
                changes_count += 1
                all_replacements.update(replacements)

        # Сохраняем документ
        doc.save(output_path)

        print(f"✓ Сохранено: {output_path}")
        print(f"  Изменено элементов: {changes_count}")

        if all_replacements:
            print(f"  Выполнено замен имен: {len(all_replacements)}")

        return output_path, all_replacements

    except Exception as e:
        print(f"✗ Ошибка: {e}")
        import traceback
        traceback.print_exc()
        return None, {}


def show_preview(input_path):
    """Предпросмотр изменений без сохранения"""
    print(f"\nПредпросмотр изменений для: {input_path}\n")

    try:
        doc = Document(input_path)
        pattern = r'(?<!\{)\{([^{}]+)\}(?!\})'

        found_tags = {}

        # Ищем в параграфах
        for paragraph in doc.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            for match in matches:
                found_tags[match] = normalize_variable_name(match)

        # Ищем в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(pattern, paragraph.text)
                        for match in matches:
                            found_tags[match] = normalize_variable_name(match)

        if found_tags:
            print(f"Найдено тегов с одинарными скобками: {len(found_tags)}")
            print("\nБудут преобразованы:")
            for original, normalized in sorted(found_tags.items()):
                if original != normalized:
                    print(f"  {{{original}}} → {{{{{normalized}}}}} ⚠️ (имя изменено)")
                else:
                    print(f"  {{{original}}} → {{{{{normalized}}}}}")
        else:
            print("Тегов с одинарными скобками не найдено.")
            print("Возможно, файл уже использует двойные скобки.")

        print()

    except Exception as e:
        print(f"✗ Ошибка при чтении файла: {e}")


if __name__ == "__main__":
    # Список файлов для обработки
    files_to_convert = [
        "ТА-турист турпродукт 27082025.docx  пр..docx",
        "для договора.docx"
    ]

    print("=" * 60)
    print("Конвертация одинарных скобок в двойные (v2)")
    print("+ Нормализация имен переменных (пробелы → подчеркивания)")
    print("=" * 60)

    # Сначала показываем предпросмотр
    for filename in files_to_convert:
        if os.path.exists(filename):
            show_preview(filename)

    # Спрашиваем подтверждение
    print("=" * 60)
    response = input("Продолжить конвертацию? (y/n): ").lower().strip()

    if response == 'y':
        print("\nНачинаем конвертацию...\n")

        all_file_replacements = {}

        for filename in files_to_convert:
            if os.path.exists(filename):
                # Создаем копию в docx_templates (перезаписываем старые)
                base_name = os.path.basename(filename)
                # Убираем лишние пробелы из имени
                clean_name = re.sub(r'\s+', ' ', base_name).strip()
                output_path = os.path.join("docx_templates", clean_name)

                result_path, replacements = convert_document(filename, output_path)
                if result_path:
                    all_file_replacements[clean_name] = replacements
            else:
                print(f"✗ Файл не найден: {filename}")

        print("\n✓ Конвертация завершена!")
        print(f"  Преобразованные файлы сохранены в: docx_templates/")

        # Показываем все замены имен
        print("\n" + "=" * 60)
        print("Сводка замен имен переменных:")
        print("=" * 60)
        for filename, replacements in all_file_replacements.items():
            if replacements:
                print(f"\n{filename}:")
                for original, normalized in sorted(replacements.items()):
                    if original != normalized:
                        print(f"  {original} → {normalized}")
    else:
        print("\nКонвертация отменена.")
