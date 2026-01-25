#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Тесты для проверки парсинга DOCX шаблонов
"""

from app import extract_template_variables
from docx import Document
import os

def test_template_parsing():
    """Тестирование парсинга всех шаблонов в папке docx_templates"""

    templates_dir = 'docx_templates'

    if not os.path.exists(templates_dir):
        print(f"❌ Директория {templates_dir} не найдена")
        return

    templates = [f for f in os.listdir(templates_dir) if f.endswith('.docx')]

    if not templates:
        print(f"❌ Шаблоны не найдены в {templates_dir}")
        return

    print(f"🔍 Найдено шаблонов: {len(templates)}\n")
    print("=" * 80)

    for template in templates:
        template_path = os.path.join(templates_dir, template)

        print(f"\n📄 Шаблон: {template}")
        print("-" * 80)

        try:
            # Парсинг переменных
            variables = extract_template_variables(template_path)

            if not variables:
                print("⚠️  Переменные не найдены!")
                continue

            print(f"✅ Найдено переменных: {len(variables)}\n")

            # Группировка по типам
            simple = [k for k, v in variables.items() if v.get('type') == 'simple']
            arrays = [k for k, v in variables.items() if v.get('type') == 'array']
            booleans = [k for k, v in variables.items() if v.get('type') == 'boolean']

            if simple:
                print("📝 Простые переменные:")
                for var in sorted(simple):
                    print(f"   - {{{{{var}}}}}")

            if arrays:
                print("\n📋 Массивы:")
                for var in sorted(arrays):
                    fields = variables[var].get('fields', [])
                    print(f"   - {var}: {fields}")

            if booleans:
                print("\n☑️  Boolean переменные:")
                for var in sorted(booleans):
                    print(f"   - {{{{{var}}}}}")

            # Проверка валидности имён переменных (Jinja2 совместимость)
            invalid_vars = []
            for var in variables.keys():
                if ' ' in var or '-' in var:
                    invalid_vars.append(var)

            if invalid_vars:
                print("\n⚠️  ВНИМАНИЕ! Найдены невалидные имена переменных:")
                for var in invalid_vars:
                    print(f"   - {{{{{var}}}}} - содержит пробелы или дефисы!")
                print("   Используйте скрипт convert_brackets_final.py для исправления")

        except Exception as e:
            print(f"❌ Ошибка при обработке: {e}")
            import traceback
            traceback.print_exc()

    print("\n" + "=" * 80)
    print("✅ Тестирование завершено!")

def test_cyrillic_support():
    """Тест поддержки кириллицы"""
    from docx import Document
    import tempfile
    import os

    print("\n🔤 Тест поддержки кириллицы")
    print("=" * 80)

    # Создаем временный документ с кириллическими переменными
    doc = Document()
    doc.add_paragraph("Договор с {{Фамилия_Имя_Отчество}}")
    doc.add_paragraph("Дата: {{дата_договора}}")
    doc.add_paragraph("Address: {{address}}")  # латиница для сравнения

    # Сохраняем временно
    temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    temp_path = temp_file.name
    doc.save(temp_path)
    temp_file.close()

    try:
        variables = extract_template_variables(temp_path)

        expected = ['Фамилия_Имя_Отчество', 'дата_договора', 'address']
        found = list(variables.keys())

        print(f"Ожидалось: {expected}")
        print(f"Найдено:   {found}")

        if set(expected) == set(found):
            print("✅ Тест пройден! Кириллица поддерживается.")
        else:
            missing = set(expected) - set(found)
            extra = set(found) - set(expected)
            if missing:
                print(f"❌ Не найдены: {missing}")
            if extra:
                print(f"⚠️  Лишние: {extra}")
    finally:
        os.unlink(temp_path)

if __name__ == '__main__':
    print("🧪 Запуск тестов для DOCX Template Filler")
    print("=" * 80)

    test_template_parsing()
    test_cyrillic_support()

    print("\n✅ Все тесты завершены!")
