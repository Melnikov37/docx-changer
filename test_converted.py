"""
Проверка сконвертированных файлов
"""
from docx import Document
import re

def check_file(filepath):
    print(f"\n{'='*60}")
    print(f"Проверка: {filepath}")
    print('='*60)

    doc = Document(filepath)

    # Паттерны для поиска
    single_pattern = r'(?<!\{)\{([^{}]+)\}(?!\})'  # Одинарные скобки
    double_pattern = r'\{\{([^{}]+)\}\}'  # Двойные скобки

    single_vars = set()
    double_vars = set()

    # Проверяем параграфы
    for para in doc.paragraphs:
        text = para.text
        single_vars.update(re.findall(single_pattern, text))
        double_vars.update(re.findall(double_pattern, text))

    # Проверяем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    single_vars.update(re.findall(single_pattern, text))
                    double_vars.update(re.findall(double_pattern, text))

    print(f"\nОдинарные скобки {{var}}: {len(single_vars)}")
    if single_vars:
        for var in sorted(single_vars):
            print(f"  {{{var}}}")

    print(f"\nДвойные скобки {{{{var}}}}: {len(double_vars)}")
    if double_vars:
        for var in sorted(double_vars):
            print(f"  {{{{{var}}}}}")

files = [
    "docx_templates/для договора.docx",
    "docx_templates/ТА-турист турпродукт 27082025.docx пр..docx"
]

for f in files:
    check_file(f)
