"""
Финальная версия конвертера для DOCX шаблонов
- Замена {var} на {{var}}
- Замена пробелов на подчеркивания
- Замена дефисов на подчеркивания (Jinja2 не поддерживает дефисы)
- Очистка множественных подчеркиваний
"""
import re
from docx import Document
import os
import json


def normalize_variable_name(name):
    """
    Нормализация имени переменной для совместимости с Jinja2
    - Заменяет пробелы на подчеркивания
    - Заменяет дефисы на подчеркивания
    - Убирает множественные подчеркивания
    """
    # Заменяем пробелы и дефисы на подчеркивания
    normalized = re.sub(r'[\s\-]+', '_', name)
    # Убираем лишние подчеркивания в начале и конце
    normalized = normalized.strip('_')
    # Заменяем множественные подчеркивания на одинарные
    normalized = re.sub(r'_+', '_', normalized)
    return normalized


def process_paragraph(paragraph):
    """Обработка параграфа - замена {var} на {{var}} и нормализация имен"""
    full_text = paragraph.text

    # Паттерн: {слово} но НЕ {{слово}} и НЕ }}
    pattern = r'(?<!\{)\{([^{}]+)\}(?!\})'

    # Проверяем, есть ли что заменять
    if not re.search(pattern, full_text):
        return False, {}

    # Словарь для отслеживания замен
    replacements = {}

    def replace_func(match):
        """Функция замены с нормализацией"""
        original_name = match.group(1)
        normalized_name = normalize_variable_name(original_name)
        replacements[original_name] = normalized_name
        return '{{' + normalized_name + '}}'

    # Заменяем одинарные скобки на двойные с нормализацией
    new_text = re.sub(pattern, replace_func, full_text)

    # Очищаем все runs в параграфе
    for run in paragraph.runs:
        run.text = ''

    # Добавляем новый текст в первый run (или создаем новый)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)

    return True, replacements


def process_table(table):
    """Обработка таблицы"""
    changed = False
    all_replacements = {}

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                was_changed, replacements = process_paragraph(paragraph)
                if was_changed:
                    changed = True
                    all_replacements.update(replacements)

    return changed, all_replacements


def convert_document(input_path, output_path=None):
    """Конвертация DOCX файла"""
    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_converted{ext}"

    print(f"Обработка: {input_path}")

    try:
        doc = Document(input_path)

        changes_count = 0
        all_replacements = {}

        # Обрабатываем параграфы
        for paragraph in doc.paragraphs:
            was_changed, replacements = process_paragraph(paragraph)
            if was_changed:
                changes_count += 1
                all_replacements.update(replacements)

        # Обрабатываем таблицы
        for table in doc.tables:
            was_changed, replacements = process_table(table)
            if was_changed:
                changes_count += 1
                all_replacements.update(replacements)

        # Сохраняем документ
        doc.save(output_path)

        print(f"✓ Сохранено: {output_path}")
        print(f"  Изменено элементов: {changes_count}")

        if all_replacements:
            print(f"  Выполнено замен имен: {len(all_replacements)}")

        return output_path, all_replacements

    except Exception as e:
        print(f"✗ Ошибка: {e}")
        import traceback
        traceback.print_exc()
        return None, {}


def generate_json_template(replacements):
    """Генерация шаблона JSON на основе найденных переменных"""
    json_template = {}
    for original, normalized in replacements.items():
        json_template[normalized] = f"<{original}>"
    return json_template


if __name__ == "__main__":
    files_to_convert = [
        "ТА-турист турпродукт 27082025.docx  пр..docx",
        "для договора.docx"
    ]

    print("=" * 70)
    print("ФИНАЛЬНАЯ конвертация шаблонов DOCX для Jinja2")
    print("Замена: пробелы → '_', дефисы → '_', {var} → {{var}}")
    print("=" * 70)

    all_file_replacements = {}

    for filename in files_to_convert:
        if os.path.exists(filename):
            # Создаем копию в docx_templates
            base_name = os.path.basename(filename)
            clean_name = re.sub(r'\s+', ' ', base_name).strip()
            output_path = os.path.join("docx_templates", clean_name)

            result_path, replacements = convert_document(filename, output_path)
            if result_path and replacements:
                all_file_replacements[clean_name] = replacements

                # Генерируем JSON шаблон
                json_template = generate_json_template(replacements)
                json_filename = os.path.join("examples",
                    clean_name.replace('.docx', '_шаблон.json'))

                with open(json_filename, 'w', encoding='utf-8') as f:
                    json.dump(json_template, f, ensure_ascii=False, indent=2)

                print(f"  JSON шаблон: {json_filename}")

        else:
            print(f"✗ Файл не найден: {filename}")

    print("\n" + "=" * 70)
    print("СВОДКА ВСЕХ ЗАМЕН:")
    print("=" * 70)

    for filename, replacements in all_file_replacements.items():
        print(f"\n📄 {filename}:")
        print("   " + "-" * 65)
        for original, normalized in sorted(replacements.items()):
            if original != normalized:
                print(f"   ❌ {{{original}}} → ✅ {{{{{normalized}}}}}")
            else:
                print(f"   ✅ {{{original}}} → {{{{{original}}}}}")

    print("\n" + "=" * 70)
    print("✓ Конвертация завершена!")
    print(f"  Шаблоны: docx_templates/")
    print(f"  JSON примеры: examples/")
    print("=" * 70)
