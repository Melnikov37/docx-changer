"""
Flask приложение для заполнения DOCX шаблонов
"""
import os
import json
import re
import uuid
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
from docxtpl import DocxTemplate
from docx import Document
from s3_client import S3Client
import db

app = Flask(__name__)

# Конфигурация
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10 MB максимум
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'output'
app.config['ALLOWED_EXTENSIONS'] = {'docx'}

# Создание необходимых директорий
for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
    Path(folder).mkdir(exist_ok=True)

# Инициализация S3 клиента и базы данных
s3_client = S3Client()
db.init_db()


def allowed_file(filename):
    """Проверка допустимого расширения файла"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def validate_file_path(base_dir, filename):
    """
    Валидация пути файла для предотвращения Path Traversal атак.

    Args:
        base_dir: Базовая директория
        filename: Имя файла

    Returns:
        Path: Безопасный путь к файлу

    Raises:
        ValueError: Если путь выходит за пределы base_dir
    """
    base_path = Path(base_dir).resolve()
    safe_filename = secure_filename(filename)
    file_path = (base_path / safe_filename).resolve()

    # Проверка, что файл находится внутри base_dir
    if not str(file_path).startswith(str(base_path)):
        raise ValueError(f"Path traversal detected: {filename}")

    return file_path


def validate_docx_file(file_stream, max_size_mb=10, max_uncompressed_mb=50):
    """
    Валидация DOCX файла для предотвращения атак (Zip Bomb, XXE).

    Args:
        file_stream: Поток байтов файла
        max_size_mb: Максимальный размер файла (сжатого)
        max_uncompressed_mb: Максимальный размер распакованного

    Returns:
        bool: True если файл валиден

    Raises:
        ValueError: Если файл не прошел валидацию
    """
    # Сохраняем текущую позицию
    original_position = file_stream.tell()

    try:
        # Проверка размера
        file_stream.seek(0, 2)  # Перейти в конец
        size = file_stream.tell()
        file_stream.seek(0)  # Вернуться в начало

        if size > max_size_mb * 1024 * 1024:
            raise ValueError(f"File too large: {size} bytes")

        # Читаем содержимое для проверки
        content = file_stream.read()
        file_stream.seek(0)

        # Проверка, что это действительно ZIP архив (DOCX)
        try:
            with zipfile.ZipFile(BytesIO(content), 'r') as zip_ref:
                # Проверка наличия обязательных файлов DOCX
                required_files = ['[Content_Types].xml', 'word/document.xml']
                file_list = zip_ref.namelist()

                for req_file in required_files:
                    if req_file not in file_list:
                        raise ValueError(f"Invalid DOCX: missing {req_file}")

                # Защита от Zip Bomb
                total_uncompressed = sum(info.file_size for info in zip_ref.infolist())
                if total_uncompressed > max_uncompressed_mb * 1024 * 1024:
                    raise ValueError(
                        f"Suspicious file: uncompressed size {total_uncompressed} exceeds limit"
                    )

                # Проверка compression ratio (защита от Zip Bomb)
                if size > 0 and total_uncompressed > size * 100:
                    raise ValueError("Suspicious compression ratio detected")

            return True

        except zipfile.BadZipFile:
            raise ValueError("File is not a valid ZIP/DOCX archive")

    finally:
        # Восстанавливаем позицию
        file_stream.seek(original_position)


def cleanup_old_files():
    """Очистка старых файлов из папок uploads и output"""
    current_time = datetime.now().timestamp()
    max_age = 3600  # 1 час

    for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
        for file_path in Path(folder).glob('*'):
            if file_path.is_file():
                file_age = current_time - file_path.stat().st_mtime
                if file_age > max_age:
                    try:
                        file_path.unlink()
                    except Exception as e:
                        app.logger.error(f"Failed to delete old file {file_path}: {e}")


def extract_template_variables(doc_path):
    """Извлечение всех Jinja2 переменных из DOCX шаблона"""
    try:
        doc = Document(doc_path)
        text_content = []

        # Извлечение текста из параграфов
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)

        # Извлечение текста из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content.append(cell.text)

        # Объединение всего текста
        full_text = ' '.join(text_content)

        # Поиск простых переменных {{variable}} (с поддержкой кириллицы) с позициями
        simple_pattern = r'\{\{\s*([a-zA-Zа-яА-ЯёЁ_][a-zA-Zа-яА-ЯёЁ0-9_\.]*?)\s*(?:\|[^}]*)?\}\}'
        simple_matches = [(m.group(1), m.start()) for m in re.finditer(simple_pattern, full_text)]

        # Поиск переменных в циклах {% for item in items %} с позициями
        loop_pattern = r'\{%\s*for\s+[a-zA-Zа-яА-ЯёЁ_][a-zA-Zа-яА-ЯёЁ0-9_]*\s+in\s+([a-zA-Zа-яА-ЯёЁ_][a-zA-Zа-яА-ЯёЁ0-9_]*)\s*%\}'
        loop_matches = [(m.group(1), m.start()) for m in re.finditer(loop_pattern, full_text)]

        # Поиск переменных в условиях {% if variable %} с позициями
        if_pattern = r'\{%\s*if\s+([a-zA-Zа-яА-ЯёЁ_][a-zA-Zа-яА-ЯёЁ0-9_]*)\s*%\}'
        if_matches = [(m.group(1), m.start()) for m in re.finditer(if_pattern, full_text)]

        # Создаем словарь позиций (первое вхождение каждой переменной)
        var_positions = {}
        for var, pos in simple_matches + loop_matches + if_matches:
            if var not in var_positions:
                var_positions[var] = pos

        # Списки переменных (для определения типов)
        simple_vars = [m[0] for m in simple_matches]
        loop_vars = [m[0] for m in loop_matches]
        if_vars = [m[0] for m in if_matches]

        # Объединение всех переменных
        all_vars = set(simple_vars + loop_vars + if_vars)

        # Разделение на простые переменные и вложенные объекты
        fields = {}
        arrays = set()

        for var in all_vars:
            # Пропускаем переменные цикла (loop.index и т.д.)
            if var.startswith('loop.'):
                continue

            # Проверяем, является ли это вложенным объектом
            if '.' in var:
                parts = var.split('.')
                root = parts[0]

                # Если корневая переменная в циклах, это массив объектов
                if root in loop_vars:
                    arrays.add(root)
                    if root not in fields:
                        fields[root] = {'type': 'array', 'fields': set(), 'position': var_positions.get(root, 9999)}
                    # Добавляем поле объекта
                    fields[root]['fields'].add(parts[1])
                else:
                    # Простой вложенный объект
                    if root not in fields:
                        fields[root] = {'type': 'object', 'fields': set(), 'position': var_positions.get(root, 9999)}
                    fields[root]['fields'].add('.'.join(parts[1:]))
            else:
                # Определяем тип переменной
                if var in loop_vars:
                    arrays.add(var)
                    if var not in fields:
                        fields[var] = {'type': 'array', 'fields': set(), 'position': var_positions.get(var, 9999)}
                elif var in if_vars:
                    # Условная переменная - вероятно boolean
                    if var not in fields:
                        fields[var] = {'type': 'boolean', 'position': var_positions.get(var, 9999)}
                else:
                    # Простая переменная
                    if var not in fields:
                        fields[var] = {'type': 'simple', 'position': var_positions.get(var, 9999)}

        # Преобразуем sets в lists для JSON
        result = {}
        for key, value in fields.items():
            if isinstance(value, dict):
                if 'fields' in value and isinstance(value['fields'], set):
                    value['fields'] = sorted(list(value['fields']))
                result[key] = value
            else:
                result[key] = value

        return result

    except Exception as e:
        app.logger.error(f"Error extracting variables: {e}")
        return {}


@app.route('/')
def index():
    """Главная страница"""
    cleanup_old_files()
    return render_template('index.html')


@app.route('/parse-template', methods=['POST'])
def parse_template():
    """Парсинг шаблона и извлечение переменных"""
    try:
        # Проверка наличия файла
        if 'template' not in request.files:
            return jsonify({'error': 'No template file provided'}), 400

        file = request.files['template']

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Only .docx files are allowed'}), 400

        # Валидация содержимого DOCX файла
        try:
            validate_docx_file(file.stream)
        except ValueError as e:
            app.logger.warning(f"File validation failed: {e}")
            return jsonify({'error': f'Invalid file: {str(e)}'}), 400

        # Сохранение временного файла
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{timestamp}_{filename}")
        file.save(temp_path)

        # Извлечение переменных
        variables = extract_template_variables(temp_path)

        # Сохраняем путь к файлу для последующего использования
        # (файл будет использован при генерации)
        session_filename = f"session_{timestamp}_{filename}"
        session_path = os.path.join(app.config['UPLOAD_FOLDER'], session_filename)
        os.rename(temp_path, session_path)

        return jsonify({
            'success': True,
            'variables': variables,
            'template_file': session_filename,
            'filename': filename
        })

    except Exception as e:
        app.logger.error(f"Error parsing template: {e}")
        # Очистка временного файла при ошибке
        if 'temp_path' in locals() and os.path.exists(temp_path):
            os.remove(temp_path)
        return jsonify({'error': f'Error parsing template: {str(e)}'}), 500


@app.route('/generate', methods=['POST'])
def generate():
    """Генерация документа из шаблона и данных JSON"""
    try:
        upload_path = None

        # Проверяем, передан ли файл напрямую или используется сохраненный
        template_file = request.form.get('template_file')

        if template_file:
            # Валидация пути для предотвращения Path Traversal
            try:
                validated_path = validate_file_path(app.config['UPLOAD_FOLDER'], template_file)
                upload_path = str(validated_path)
            except ValueError as e:
                app.logger.warning(f"Path traversal attempt: {e}")
                return jsonify({'error': 'Invalid file path'}), 400

            if not os.path.exists(upload_path):
                return jsonify({'error': 'Template file not found. Please upload again.'}), 400
            # Извлекаем оригинальное имя файла
            filename = template_file.replace('session_', '').split('_', 2)[2] if '_' in template_file else template_file
        else:
            # Проверка наличия файла в запросе
            if 'template' not in request.files:
                return jsonify({'error': 'No template file provided'}), 400

            file = request.files['template']

            if file.filename == '':
                return jsonify({'error': 'No file selected'}), 400

            if not allowed_file(file.filename):
                return jsonify({'error': 'Only .docx files are allowed'}), 400

            # Валидация содержимого DOCX файла
            try:
                validate_docx_file(file.stream)
            except ValueError as e:
                app.logger.warning(f"File validation failed: {e}")
                return jsonify({'error': f'Invalid file: {str(e)}'}), 400

            # Сохранение загруженного шаблона
            filename = secure_filename(file.filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            upload_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{timestamp}_{filename}")
            file.save(upload_path)

        # Получение JSON данных
        try:
            json_data = request.form.get('data', '{}')
            context = json.loads(json_data)
        except json.JSONDecodeError as e:
            return jsonify({'error': f'Invalid JSON: {str(e)}'}), 400

        # Генерация имени выходного файла
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"filled_{timestamp}_{filename}"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

        # Обработка шаблона
        try:
            doc = DocxTemplate(upload_path)
            doc.render(context)
            doc.save(output_path)
        except Exception as e:
            return jsonify({'error': f'Template processing error: {str(e)}'}), 500

        # Сохранение в историю (MinIO + БД)
        try:
            # Генерируем уникальный ключ для S3
            s3_key = f"generated/{uuid.uuid4()}_{output_filename}"

            # Загружаем в MinIO
            if s3_client.upload_file(output_path, s3_key):
                # Получаем размер файла
                file_size = os.path.getsize(output_path)

                # Сохраняем метаданные в БД
                doc_id = db.add_generated_document(
                    template_name=filename,
                    output_filename=output_filename,
                    s3_key=s3_key,
                    json_data=context,
                    file_size=file_size
                )

                app.logger.info(f"Document saved to history: ID={doc_id}, S3={s3_key}")
        except Exception as e:
            app.logger.error(f"Error saving to history: {e}")
            # Продолжаем даже если сохранение в историю не удалось

        return jsonify({
            'success': True,
            'filename': output_filename,
            'message': 'Document generated successfully'
        })

    except Exception as e:
        app.logger.error(f"Error in generate endpoint: {e}")
        return jsonify({'error': f'Server error: {str(e)}'}), 500


@app.route('/download/<filename>')
def download(filename):
    """Скачивание сгенерированного документа"""
    try:
        # Валидация пути для предотвращения Path Traversal
        try:
            validated_path = validate_file_path(app.config['OUTPUT_FOLDER'], filename)
            file_path = str(validated_path)
        except ValueError as e:
            app.logger.warning(f"Path traversal attempt in download: {e}")
            return jsonify({'error': 'Invalid file path'}), 400

        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404

        return send_file(
            file_path,
            as_attachment=True,
            download_name=secure_filename(filename),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        app.logger.error(f"Error in download endpoint: {e}")
        return jsonify({'error': f'Download error: {str(e)}'}), 500


@app.route('/health')
def health():
    """Проверка состояния приложения"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat()
    })


@app.route('/templates', methods=['GET'])
def get_templates():
    """Получение списка всех шаблонов"""
    try:
        templates = db.get_all_templates()
        return jsonify({'success': True, 'templates': templates})
    except Exception as e:
        app.logger.error(f"Error getting templates: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/templates/save', methods=['POST'])
def save_template():
    """Сохранение шаблона в библиотеку"""
    try:
        template_file = request.form.get('template_file')
        name = request.form.get('name')
        description = request.form.get('description', '')

        if not template_file or not name:
            return jsonify({'error': 'Template file and name are required'}), 400

        # Валидация пути для предотвращения Path Traversal
        try:
            validated_path = validate_file_path(app.config['UPLOAD_FOLDER'], template_file)
            upload_path = str(validated_path)
        except ValueError as e:
            app.logger.warning(f"Path traversal attempt in save_template: {e}")
            return jsonify({'error': 'Invalid file path'}), 400

        if not os.path.exists(upload_path):
            return jsonify({'error': 'Template file not found'}), 404

        # Извлечение переменных из шаблона
        variables = extract_template_variables(upload_path)

        # Генерация уникального ключа для S3
        s3_key = f"{uuid.uuid4()}_{secure_filename(template_file)}"

        # Загрузка в S3
        if not s3_client.upload_file(upload_path, s3_key):
            return jsonify({'error': 'Failed to upload to storage'}), 500

        # Сохранение метаданных в БД
        template_id = db.add_template(
            name=name,
            original_filename=template_file,
            s3_key=s3_key,
            description=description,
            variables=variables
        )

        return jsonify({
            'success': True,
            'template_id': template_id,
            'message': 'Template saved successfully'
        })

    except Exception as e:
        app.logger.error(f"Error saving template: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/templates/<int:template_id>', methods=['GET'])
def load_template(template_id):
    """Загрузка шаблона из библиотеки"""
    try:
        template = db.get_template(template_id)

        if not template:
            return jsonify({'error': 'Template not found'}), 404

        # Скачиваем файл из S3 во временную папку
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        temp_filename = f"session_{timestamp}_{template['original_filename']}"
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_filename)

        if not s3_client.download_file(template['s3_key'], temp_path):
            return jsonify({'error': 'Failed to load template from storage'}), 500

        # Парсим переменные
        variables = extract_template_variables(temp_path)

        return jsonify({
            'success': True,
            'template_file': temp_filename,
            'variables': variables,
            'name': template['name'],
            'description': template['description']
        })

    except Exception as e:
        app.logger.error(f"Error loading template: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/templates/<int:template_id>', methods=['DELETE'])
def delete_template_endpoint(template_id):
    """Удаление шаблона из библиотеки"""
    try:
        s3_key = db.delete_template(template_id)

        if not s3_key:
            return jsonify({'error': 'Template not found'}), 404

        # Удаляем файл из S3
        s3_client.delete_file(s3_key)

        return jsonify({'success': True, 'message': 'Template deleted successfully'})

    except Exception as e:
        app.logger.error(f"Error deleting template: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/history', methods=['GET'])
def get_history():
    """Получение истории сгенерированных документов"""
    try:
        limit = request.args.get('limit', 100, type=int)
        documents = db.get_all_generated_documents(limit=limit)
        return jsonify({'success': True, 'documents': documents})
    except Exception as e:
        app.logger.error(f"Error getting history: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/history/<int:doc_id>/download', methods=['GET'])
def download_from_history(doc_id):
    """Скачивание документа из истории"""
    try:
        document = db.get_generated_document(doc_id)

        if not document:
            return jsonify({'error': 'Document not found'}), 404

        # Скачиваем файл из S3 во временную папку
        temp_path = os.path.join(app.config['OUTPUT_FOLDER'], document['output_filename'])

        if not s3_client.download_file(document['s3_key'], temp_path):
            return jsonify({'error': 'Failed to download from storage'}), 500

        return send_file(
            temp_path,
            as_attachment=True,
            download_name=document['output_filename'],
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        app.logger.error(f"Error downloading from history: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/history/<int:doc_id>/data', methods=['GET'])
def get_history_data(doc_id):
    """Получение JSON данных документа из истории"""
    try:
        document = db.get_generated_document(doc_id)

        if not document:
            return jsonify({'error': 'Document not found'}), 404

        return jsonify({
            'success': True,
            'template_name': document['template_name'],
            'output_filename': document['output_filename'],
            'json_data': document.get('json_data', {}),
            'created_at': document['created_at']
        })

    except Exception as e:
        app.logger.error(f"Error getting document data: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/history/<int:doc_id>', methods=['DELETE'])
def delete_from_history(doc_id):
    """Удаление документа из истории"""
    try:
        s3_key = db.delete_generated_document(doc_id)

        if not s3_key:
            return jsonify({'error': 'Document not found'}), 404

        # Удаляем файл из S3
        s3_client.delete_file(s3_key)

        return jsonify({'success': True, 'message': 'Document deleted successfully'})

    except Exception as e:
        app.logger.error(f"Error deleting from history: {e}")
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True, host='127.0.0.1', port=5001)
