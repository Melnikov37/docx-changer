"""
Скрипт для конвертации одинарных фигурных скобок {variable}
в двойные {{variable}} для совместимости с Jinja2
"""
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sys
import os


def process_paragraph(paragraph):
    """Обработка параграфа - замена {var} на {{var}}"""
    # Собираем весь текст параграфа
    full_text = paragraph.text

    # Находим все теги с одинарными скобками
    # Паттерн: {слово} но НЕ {{слово}} и НЕ }}
    pattern = r'(?<!\{)\{([^{}]+)\}(?!\})'

    # Проверяем, есть ли что заменять
    if not re.search(pattern, full_text):
        return False

    # Заменяем одинарные скобки на двойные
    new_text = re.sub(pattern, r'{{\1}}', full_text)

    # Очищаем все runs в параграфе
    for run in paragraph.runs:
        run.text = ''

    # Добавляем новый текст в первый run (или создаем новый)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)

    return True


def process_table(table):
    """Обработка таблицы"""
    changed = False
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if process_paragraph(paragraph):
                    changed = True
    return changed


def convert_document(input_path, output_path=None):
    """Конвертация DOCX файла"""
    if output_path is None:
        # Создаем имя выходного файла
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_converted{ext}"

    print(f"Обработка: {input_path}")

    try:
        # Загружаем документ
        doc = Document(input_path)

        changes_count = 0

        # Обрабатываем параграфы
        for paragraph in doc.paragraphs:
            if process_paragraph(paragraph):
                changes_count += 1

        # Обрабатываем таблицы
        for table in doc.tables:
            if process_table(table):
                changes_count += 1

        # Сохраняем документ
        doc.save(output_path)

        print(f"✓ Сохранено: {output_path}")
        print(f"  Изменено элементов: {changes_count}")

        return output_path

    except Exception as e:
        print(f"✗ Ошибка: {e}")
        return None


def show_preview(input_path):
    """Предпросмотр изменений без сохранения"""
    print(f"\nПредпросмотр изменений для: {input_path}\n")

    try:
        doc = Document(input_path)
        pattern = r'(?<!\{)\{([^{}]+)\}(?!\})'

        found_tags = set()

        # Ищем в параграфах
        for paragraph in doc.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            found_tags.update(matches)

        # Ищем в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(pattern, paragraph.text)
                        found_tags.update(matches)

        if found_tags:
            print(f"Найдено тегов с одинарными скобками: {len(found_tags)}")
            print("\nБудут преобразованы:")
            for tag in sorted(found_tags):
                print(f"  {{{tag}}} → {{{{{tag}}}}}")
        else:
            print("Тегов с одинарными скобками не найдено.")
            print("Возможно, файл уже использует двойные скобки.")

        print()

    except Exception as e:
        print(f"✗ Ошибка при чтении файла: {e}")


if __name__ == "__main__":
    # Список файлов для обработки
    files_to_convert = [
        "ТА-турист турпродукт 27082025.docx  пр..docx",
        "для договора.docx"
    ]

    print("=" * 60)
    print("Конвертация одинарных скобок в двойные")
    print("=" * 60)

    # Сначала показываем предпросмотр
    for filename in files_to_convert:
        if os.path.exists(filename):
            show_preview(filename)

    # Спрашиваем подтверждение
    print("=" * 60)
    response = input("Продолжить конвертацию? (y/n): ").lower().strip()

    if response == 'y':
        print("\nНачинаем конвертацию...\n")

        for filename in files_to_convert:
            if os.path.exists(filename):
                # Создаем копию в docx_templates
                base_name = os.path.basename(filename)
                # Убираем лишние пробелы из имени
                clean_name = re.sub(r'\s+', ' ', base_name).strip()
                output_path = os.path.join("docx_templates", clean_name)

                convert_document(filename, output_path)
            else:
                print(f"✗ Файл не найден: {filename}")

        print("\n✓ Конвертация завершена!")
        print(f"  Преобразованные файлы сохранены в: docx_templates/")
    else:
        print("\nКонвертация отменена.")
